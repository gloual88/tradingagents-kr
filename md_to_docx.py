"""Convert TradingAgents complete_report.md to a Korean-friendly .docx.

Usage:
    python md_to_docx.py SPY 2026-05-15
    python md_to_docx.py SPY 2026-05-15 --out d:/path/to/output.docx
"""
from __future__ import annotations

import argparse
import os
import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


KOR_FONT = "Malgun Gothic"
CODE_FONT = "Consolas"


def _set_font(run, name=KOR_FONT, size=11, bold=False, italic=False,
              color=None, mono=False):
    run.font.name = name if not mono else CODE_FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        from docx.oxml import OxmlElement
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), KOR_FONT)
    rfonts.set(qn("w:ascii"), CODE_FONT if mono else name)
    rfonts.set(qn("w:hAnsi"), CODE_FONT if mono else name)


INLINE_RE = re.compile(
    r"(\*\*(?P<bold>[^*]+)\*\*)"
    r"|(\*(?P<italic>[^*]+)\*)"
    r"|(`(?P<code>[^`]+)`)"
)


def _add_inline(paragraph, text: str, base_bold=False, base_size=11):
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
            _set_font(run, bold=base_bold, size=base_size)
        if m.group("bold") is not None:
            run = paragraph.add_run(m.group("bold"))
            _set_font(run, bold=True, size=base_size)
        elif m.group("italic") is not None:
            run = paragraph.add_run(m.group("italic"))
            _set_font(run, italic=True, size=base_size)
        elif m.group("code") is not None:
            run = paragraph.add_run(m.group("code"))
            _set_font(run, mono=True, size=base_size - 1,
                      color=RGBColor(0xB0, 0x30, 0x30))
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        _set_font(run, bold=base_bold, size=base_size)


def _add_heading(doc, text: str, level: int):
    sizes = {1: 18, 2: 15, 3: 13, 4: 12}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    _set_font(run, size=sizes.get(level, 11), bold=True,
              color=RGBColor(0x1F, 0x3A, 0x68))


def _add_table(doc, rows: list[list[str]]):
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Light Grid Accent 1"
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            _add_inline(p, cell_text, base_bold=(r_idx == 0), base_size=10)


def md_to_docx(md_text: str, out_path: Path, title: str | None = None):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = KOR_FONT
    style.font.size = Pt(11)

    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        _set_font(run, size=22, bold=True, color=RGBColor(0x0D, 0x1E, 0x40))

    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if not line.strip():
            i += 1
            continue

        if line.startswith("```"):
            i += 1
            buf = []
            while i < len(lines) and not lines[i].startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(12)
            run = p.add_run("\n".join(buf))
            _set_font(run, mono=True, size=10,
                      color=RGBColor(0x33, 0x33, 0x33))
            continue

        m = re.match(r"^(#{1,4})\s+(.*)$", line)
        if m:
            _add_heading(doc, m.group(2).strip(), len(m.group(1)))
            i += 1
            continue

        if re.match(r"^[-*]{3,}\s*$", line):
            doc.add_paragraph("─" * 60)
            i += 1
            continue

        if "|" in line and i + 1 < len(lines) and re.match(
            r"^\s*\|?[\s:|-]+\|[\s:|-]+", lines[i + 1]
        ):
            rows = []
            while i < len(lines) and "|" in lines[i]:
                row_line = lines[i].strip().strip("|")
                if re.match(r"^[\s:|-]+$", row_line):
                    i += 1
                    continue
                cells = [c.strip() for c in row_line.split("|")]
                rows.append(cells)
                i += 1
            _add_table(doc, rows)
            continue

        m = re.match(r"^(\s*)[-*+]\s+(.*)$", line)
        if m:
            depth = len(m.group(1)) // 2
            style_name = "List Bullet" if depth == 0 else "List Bullet 2"
            try:
                p = doc.add_paragraph(style=style_name)
            except KeyError:
                p = doc.add_paragraph()
            _add_inline(p, m.group(2))
            i += 1
            continue

        m = re.match(r"^\s*(\d+)\.\s+(.*)$", line)
        if m:
            try:
                p = doc.add_paragraph(style="List Number")
            except KeyError:
                p = doc.add_paragraph()
            _add_inline(p, m.group(2))
            i += 1
            continue

        if line.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            _add_inline(p, line.lstrip(">").strip())
            i += 1
            continue

        p = doc.add_paragraph()
        _add_inline(p, line)
        i += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


REPORT_ORDER = [
    ("market_report.md",            "기술적 분석 (Market Analyst)"),
    ("sentiment_report.md",         "투자심리 분석 (Sentiment Analyst)"),
    ("news_report.md",              "뉴스 분석 (News Analyst)"),
    ("fundamentals_report.md",      "펀더멘털 분석 (Fundamentals Analyst)"),
    ("investment_plan.md",          "리서치팀 토론 결과 (Bull vs Bear)"),
    ("trader_investment_plan.md",   "트레이더 투자 계획"),
    ("final_trade_decision.md",     "최종 매매 결정 (Risk Management)"),
]


def _assemble_from_reports(reports_dir: Path) -> str:
    parts = []
    for fname, section_title in REPORT_ORDER:
        f = reports_dir / fname
        if not f.exists():
            continue
        body = f.read_text(encoding="utf-8").strip()
        parts.append(f"# {section_title}\n\n{body}")
    return "\n\n---\n\n".join(parts)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("ticker")
    parser.add_argument("date", help="YYYY-MM-DD")
    parser.add_argument("--logs-root",
                        default=os.path.expanduser("~/.tradingagents/logs"))
    parser.add_argument("--out", default=None)
    args = parser.parse_args()

    run_dir = Path(args.logs_root) / args.ticker / args.date
    complete_md = run_dir / "complete_report.md"
    reports_dir = run_dir / "reports"

    if complete_md.exists():
        md_text = complete_md.read_text(encoding="utf-8")
        source_label = "complete_report.md"
    elif reports_dir.exists():
        md_text = _assemble_from_reports(reports_dir)
        if not md_text:
            raise SystemExit(f"리포트가 비어 있습니다: {reports_dir}")
        source_label = f"reports/ ({len(REPORT_ORDER)}개 파일 조립)"
    else:
        raise SystemExit(f"리포트 폴더를 찾을 수 없습니다: {run_dir}")

    reports_out = Path(__file__).parent / "reports"
    reports_out.mkdir(parents=True, exist_ok=True)
    default_out = reports_out / f"{args.ticker}_{args.date}.docx"
    out_path = Path(args.out) if args.out else default_out
    title = f"TradingAgents 분석 리포트 — {args.ticker} ({args.date})"
    md_to_docx(md_text, out_path, title=title)
    print(f"소스: {source_label}")
    print(f"저장 완료: {out_path}")


if __name__ == "__main__":
    main()
